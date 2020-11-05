import itertools
import os
import re
from pathlib import Path
from typing import List, Union

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Length
from docx.table import Table
from lxml import etree

from app_path import application_path
from cif.cif_file_io import CifContainer
from cif.spgrps import SpaceGroups
from cif.text import retranslate_delimiter
from tools import grouper

# protected space character:
prot_space = u'\u00A0'
# Angstrom character:
angstrom = u'\u212B'
# bigger or equal:
bequal = u'\u2265'
# small_sigma:
sigma_sm = u'\u03C3'
# en dash:
halbgeviert = u'\u2013'
# degree sign:
degree_sign = u'\u00B0'
# middle ellipsis
ellipsis_mid = u'\u22EF'
# ellipsis
ellipsis = u'\u2026'
# less or equal sign
lessequal = u'\u2264'
# times (cross) symbol
timessym = u'\u00d7'
# lambda
lambdasym = u'\u03bb'
# one bar
one_bar = u'\u0031\u0305'
# Zero with space ZWSP
zero_width_space = u'\u200B'

cif_keywords_list = (
    ['_chemical_formula_weight', 1],
    ['_diffrn_ambient_temperature', 2],
    ['_space_group_crystal_system', 3],
    ['_cell_length_a', 5],
    ['_cell_length_b', 6],
    ['_cell_length_c', 7],
    ['_cell_angle_alpha', 8],
    ['_cell_angle_beta', 9],
    ['_cell_angle_gamma', 10],
    ['_cell_volume', 11],
    ['_cell_formula_units_Z', 12],
    ['_exptl_crystal_density_diffrn', 13],
    ['_exptl_absorpt_coefficient_mu', 14],
    ['_exptl_crystal_F_000', 15],
    ['_exptl_crystal_colour', 17],
    ['_exptl_crystal_description', 18],
    ['_diffrn_reflns_theta_min', 20],
    ['_diffrn_reflns_theta_max', 20],
    ['_diffrn_reflns_number', 22],
    ['_refine_ls_goodness_of_fit_ref', 25],

)


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def this_or_quest(value):
    """
    Returns the value or a question mark if the value is None.
    """
    return value if value else '?'


def create_document(report_docx_path: str) -> Document:
    """
    Creates the report docx document.
    :param report_docx_path: Path to the report file.
    :return: The document instance.
    """
    try:
        document = Document(Path(report_docx_path).joinpath(application_path, 'template/template1.docx').absolute())
    except FileNotFoundError as e:
        print(e)
        document = Document()
    # Deleting first (empty) paragraph, otherwise first line would be an empty one:
    try:
        p = document.paragraphs[0]
        delete_paragraph(p)
    except IndexError:
        # no paragraph there
        pass
    return document


def set_column_width(column, width: Length) -> None:
    for cell in column.cells:
        cell.width = width


def isfloat(value: Union[str, int, float]) -> bool:
    try:
        float(value)
        return True
    except ValueError:
        return False


def math_to_word(eq: str) -> str:
    """Transform a sympy equation to be printed in word document."""
    tree = etree.fromstring(eq)
    xslt = etree.parse(os.path.join(application_path, 'template/mathml2omml.xsl'))
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()


def format_radiation(radiation_type: str) -> list:
    radtype = list(radiation_type.partition("K"))
    if len(radtype) > 2:
        radtype[2] = retranslate_delimiter(radtype[2])
        return radtype
    else:
        return radtype


def format_space_group(table: Table, cif: CifContainer) -> None:
    """
    Sets formating of the space group symbol in row 6.
    """
    space_group = cif['_space_group_name_H-M_alt'].strip("'")
    it_number = cif['_space_group_IT_number']
    paragraph = table.cell(5, 1).paragraphs[0]
    try:
        # The HM space group symbol
        s = SpaceGroups()
        spgrxml = s.iucrNumberToMathml(it_number)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph._element.append(math_to_word(spgrxml))
        paragraph.add_run(' (' + it_number + ')')
    except Exception:
        # Use fallback:
        if space_group:
            if len(space_group) > 4:  # don't modify P 1
                space_group = re.sub(r'\s1', '', space_group)  # remove extra Hall "1" for mono and tric
            space_group = re.sub(r'\s', '', space_group)  # remove all remaining whitespace
            # space_group = re.sub(r'-1', one_bar, space_group)  # exchange -1 with 1bar
            space_group_formated_text = [char for char in space_group]  # ???)
            is_sub = False
            for k, char in enumerate(space_group_formated_text):
                sgrunsub = paragraph.add_run(char)
                if not char.isdigit():
                    sgrunsub.font.italic = True
                else:
                    if space_group_formated_text[k - 1].isdigit() and not is_sub:
                        is_sub = True
                        sgrunsub.font.subscript = True  # lowercase the second digit if previous is also digit
                    else:
                        is_sub = False  # only every second number as subscript for P212121 etc.
            if it_number:
                paragraph.add_run(' (' + it_number + ')')
        else:
            paragraph.add_run('?')


def populate_description_columns(main_table: Table, cif: CifContainer) -> None:
    """
    This Method adds the descriptions to the fist property table column.
    """
    lgnd1 = main_table.cell(1, 0).paragraphs[0].add_run('Empirical formula')
    lgnd2 = main_table.cell(2, 0).paragraphs[0].add_run('Formula weight')
    lgnd3 = main_table.cell(3, 0).paragraphs[0].add_run('Temperature [K]')
    lgnd4 = main_table.cell(4, 0).paragraphs[0].add_run('Crystal system')
    lgnd5 = main_table.cell(5, 0).paragraphs[0].add_run('Space group (number)')
    lgnd6 = main_table.cell(6, 0).paragraphs[0]
    lgnd6.add_run('a').font.italic = True
    lgnd6.add_run(' [{}]'.format(angstrom))
    lgnd7 = main_table.cell(7, 0).paragraphs[0]
    lgnd7.add_run('b').font.italic = True
    lgnd7.add_run(' [{}]'.format(angstrom))
    lgnd8 = main_table.cell(8, 0).paragraphs[0]
    lgnd8.add_run('c').font.italic = True
    lgnd8.add_run(' [{}]'.format(angstrom))
    lgnd9 = main_table.cell(9, 0).paragraphs[0].add_run('\u03B1 [{}]'.format(angstrom))
    lgnd10 = main_table.cell(10, 0).paragraphs[0].add_run('\u03B2 [{}]'.format(angstrom))
    lgnd11 = main_table.cell(11, 0).paragraphs[0].add_run('\u03B3 [{}]'.format(angstrom))
    lgnd12 = main_table.cell(12, 0).paragraphs[0]
    lgnd12.add_run('Volume [{}'.format(angstrom))
    lgnd12.add_run('3').font.superscript = True
    lgnd12.add_run(']')
    lgnd13 = main_table.cell(13, 0).paragraphs[0].add_run('Z').font.italic = True
    lgnd14 = main_table.cell(14, 0).paragraphs[0]
    lgnd14.add_run('\u03C1').font.italic = True
    lgnd14.add_run('calc').font.subscript = True
    lgnd14.add_run(' [g/cm')
    lgnd14.add_run('3').font.superscript = True
    lgnd14.add_run(']')
    lgnd15 = main_table.cell(15, 0).paragraphs[0]
    lgnd15.add_run('\u03BC').font.italic = True
    lgnd15.add_run(' [mm')
    lgnd15.add_run('-1').font.superscript = True
    lgnd15.add_run(']')
    lgnd16 = main_table.cell(16, 0).paragraphs[0]
    lgnd16.add_run('F').font.italic = True
    lgnd16.add_run('(000)')
    lgnd17 = main_table.cell(17, 0).paragraphs[0]
    lgnd17.add_run('Crystal size [mm')
    lgnd17.add_run('3').font.superscript = True
    lgnd17.add_run(']')
    lgnd18 = main_table.cell(18, 0).paragraphs[0].add_run('Crystal colour')
    lgnd19 = main_table.cell(19, 0).paragraphs[0].add_run('Crystal shape')
    lgnd20 = main_table.cell(20, 0).paragraphs[0].add_run('Radiation')
    lgnd21 = main_table.cell(21, 0).paragraphs[0].add_run('2\u03F4 range [\u00b0]')
    lgnd22 = main_table.cell(22, 0).paragraphs[0].add_run('Index ranges')
    lgnd23 = main_table.cell(23, 0).paragraphs[0].add_run('Reflections collected')
    lgnd24 = main_table.cell(24, 0).paragraphs[0].add_run('Independent reflections')
    lgnd25 = main_table.cell(25, 0).paragraphs[0]
    theta_full = cif['_diffrn_reflns_theta_full']
    if theta_full:
        lgnd25.add_run('Completeness to \n\u03B8 = {}°'.format(theta_full))
    else:
        lgnd25.add_run('Completeness')
    main_table.cell(26, 0).paragraphs[0].add_run('Data / Restraints / Parameters')
    lgnd27 = main_table.cell(27, 0).paragraphs[0]
    lgnd27.add_run('Goodness-of-fit on ')
    lgnd27.add_run('F').font.italic = True
    lgnd27.add_run('2').font.superscript = True
    lgnd28 = main_table.cell(28, 0).paragraphs[0]
    lgnd28.add_run('Final ')
    lgnd28.add_run('R').font.italic = True
    lgnd28.add_run(' indexes \n[')
    lgnd28.add_run('I').font.italic = True
    lgnd28.add_run('{}2{}('.format(bequal, sigma_sm))
    lgnd28.add_run('I').font.italic = True
    lgnd28.add_run(')]')
    lgnd29 = main_table.cell(29, 0).paragraphs[0]
    lgnd29.add_run('Final ')
    lgnd29.add_run('R').font.italic = True
    lgnd29.add_run(' indexes \n[all data]')
    lgnd30 = main_table.cell(30, 0).paragraphs[0]
    lgnd30.add_run('Largest peak/hole [e{}'.format(angstrom))
    lgnd30.add_run('3').font.superscript = True
    lgnd30.add_run(']')
    if not cif.is_centrosymm:
        lgnd31 = main_table.cell(31, 0).paragraphs[0]
        lgnd31.add_run('Flack X parameter')
    exti = cif['_refine_ls_extinction_coef']
    if exti not in ['.', "'.'", '?', '']:
        # always the last cell
        num = len(main_table.columns[0].cells)
        main_table.columns[0].cells[num - 1].paragraphs[0].add_run('Extinction coefficient')


def populate_main_table_values(main_table: Table, cif: CifContainer):
    """
    Fills the main table with residuals. Column, by column.
    """
    header_cells = main_table.rows[0].cells
    header_cells[0].paragraphs[0].add_run('CCDC number')  # .bold = True
    header_cells[1].paragraphs[0].add_run(cif['_database_code_depnum_ccdc_archive'])  # .bold = True

    # Set text for all usual cif keywords by a lookup table:
    for _, key in enumerate(cif_keywords_list):
        # key[1] contains the row number:
        cell = main_table.cell(key[1] + 1, 1)
        if cif[key[0]]:
            cell.text = cif[key[0]]
        else:
            cell.text = '?'
            continue
    # Now the special handling:
    # The sum formula:
    sum_formula = 'no sum formula'
    if cif['_chemical_formula_sum']:
        sum_formula = cif['_chemical_formula_sum']
        ltext2 = sum_formula.replace(" ", "").replace("'", "")
        ltext3 = [''.join(x[1]) for x in itertools.groupby(ltext2, lambda x: x.isalpha())]
        for _, word in enumerate(ltext3):
            formrun = main_table.cell(1, 1).paragraphs[0]
            formrunsub = formrun.add_run(word)
            if isfloat(word):
                formrunsub.font.subscript = True

    format_space_group(main_table, cif)
    radiation_type = cif['_diffrn_radiation_type']
    radiation_wavelength = cif['_diffrn_radiation_wavelength']
    crystal_size_min = cif['_exptl_crystal_size_min']
    crystal_size_mid = cif['_exptl_crystal_size_mid']
    crystal_size_max = cif['_exptl_crystal_size_max']
    limit_h_min = cif['_diffrn_reflns_limit_h_min']
    limit_h_max = cif['_diffrn_reflns_limit_h_max']
    limit_k_min = cif['_diffrn_reflns_limit_k_min']
    limit_k_max = cif['_diffrn_reflns_limit_k_max']
    theta_min = cif['_diffrn_reflns_theta_min']
    theta_max = cif['_diffrn_reflns_theta_max']
    limit_l_min = cif['_diffrn_reflns_limit_l_min']
    limit_l_max = cif['_diffrn_reflns_limit_l_max']
    reflns_number_total = cif['_reflns_number_total']
    reflns_av_R_equivalents = cif['_diffrn_reflns_av_R_equivalents']
    reflns_av_unetI = cif['_diffrn_reflns_av_unetI/netI']
    ls_number_reflns = cif['_refine_ls_number_reflns']
    ls_number_restraints = cif['_refine_ls_number_restraints']
    ls_number_parameters = cif['_refine_ls_number_parameters']
    ls_R_factor_gt = cif['_refine_ls_R_factor_gt']
    ls_wR_factor_gt = cif['_refine_ls_wR_factor_gt']
    ls_R_factor_all = cif['_refine_ls_R_factor_all']
    ls_wR_factor_ref = cif['_refine_ls_wR_factor_ref']
    goof = cif['_refine_ls_goodness_of_fit_ref']
    try:
        completeness = "{0:.1f} %".format(round(float(cif['_diffrn_measured_fraction_theta_full']) * 100, 1))
    except ValueError:
        completeness = '?'
    try:
        diff_density_min = "{0:.2f}".format(round(float(cif['_refine_diff_density_min']), 2))
    except ValueError:
        diff_density_min = '?'
    try:
        diff_density_max = "{0:.2f}".format(round(float(cif['_refine_diff_density_max']), 2))
    except ValueError:
        diff_density_max = '?'

    # now prepare & write all the concatenated & derived cell contents:
    main_table.cell(17, 1).text = this_or_quest(crystal_size_max) + timessym + \
                                  this_or_quest(crystal_size_mid) + timessym + \
                                  this_or_quest(crystal_size_min)
    wavelength = str(' ({} ='.format(lambdasym) + this_or_quest(radiation_wavelength) +
                     '{}{})'.format(prot_space, angstrom)).replace(' ', '')
    # radtype: ('Mo', 'K', '\\a')
    radtype = format_radiation(radiation_type)
    radrun = main_table.cell(20, 1).paragraphs[0]
    # radiation type e.g. Mo:
    radrun.add_run(radtype[0])
    # K line:
    radrunita = radrun.add_run(radtype[1])
    radrunita.font.italic = True
    alpha = radrun.add_run(radtype[2])
    alpha.font.italic = True
    alpha.font.subscript = True
    # wavelength lambda:
    radrun.add_run(' ' + wavelength)
    try:
        d_max = ' ({:.2f}{}{})'.format(float(radiation_wavelength) / (2 * sin(radians(float(theta_max)))), prot_space,
                                       angstrom)
        # 2theta range:
        main_table.cell(21, 1).text = "{:.2f} to {:.2f}{}".format(2 * float(theta_min), 2 * float(theta_max), d_max)
    except ValueError:
        main_table.cell(21, 1).text = '? to ?'
    main_table.cell(22, 1).text = limit_h_min + ' {} h {} '.format(lessequal, lessequal) + limit_h_max + '\n' \
                                  + limit_k_min + ' {} k {} '.format(lessequal, lessequal) + limit_k_max + '\n' \
                                  + limit_l_min + ' {} l {} '.format(lessequal, lessequal) + limit_l_max
    rint_p = main_table.cell(24, 1).paragraphs[0]
    rint_p.add_run(this_or_quest(reflns_number_total) + '\n')
    rint_p.add_run('R').font.italic = True
    rint_p.add_run('int').font.subscript = True
    rint_p.add_run(' = ' + this_or_quest(reflns_av_R_equivalents) + '\n')
    rint_p.add_run('R').font.italic = True
    rint_p.add_run('sigma').font.subscript = True
    rint_p.add_run(' = ' + this_or_quest(reflns_av_unetI))
    main_table.cell(25, 1).paragraphs[0].add_run(completeness)
    main_table.cell(26, 1).text = this_or_quest(ls_number_reflns) + '/' \
                                  + this_or_quest(ls_number_restraints) + '/' \
                                  + this_or_quest(ls_number_parameters)
    main_table.cell(27, 1).paragraphs[0].add_run(goof)
    r2sig_p = main_table.cell(28, 1).paragraphs[0]
    r2sig_p.add_run('R').font.italic = True
    r2sig_p.add_run('1').font.subscript = True
    r2sig_p.add_run(' = ' + this_or_quest(ls_R_factor_gt))
    r2sig_p.add_run('\nw')
    r2sig_p.add_run('R').font.italic = True
    r2sig_p.add_run('2').font.subscript = True
    r2sig_p.add_run(' = ' + this_or_quest(ls_wR_factor_gt))
    rfull_p = main_table.cell(29, 1).paragraphs[0]
    rfull_p.add_run('R').font.italic = True
    rfull_p.add_run('1').font.subscript = True
    rfull_p.add_run(' = ' + this_or_quest(ls_R_factor_all))
    rfull_p.add_run('\nw')
    rfull_p.add_run('R').font.italic = True
    rfull_p.add_run('2').font.subscript = True
    rfull_p.add_run(' = ' + ls_wR_factor_ref)
    main_table.cell(30, 1).text = diff_density_max + '/' + diff_density_min
    if not cif.is_centrosymm:
        main_table.cell(31, 1).text = cif['_refine_ls_abs_structure_Flack'] or '?'
    exti = cif['_refine_ls_extinction_coef']
    if exti not in ['.', "'.'", '?', '']:
        num = len(main_table.columns[0].cells)
        main_table.columns[1].cells[num - 1].text = exti


def add_residuals_table(document: Document(), cif: CifContainer, table_num: int) -> int:
    # table_num += 1
    exti = cif['_refine_ls_extinction_coef']
    rows = 33
    if cif.is_centrosymm:
        rows -= 1
    # Remove one row for the extinction coefficient:
    if exti in ['.', "'.'", '?', '']:
        rows -= 1
    main_table = document.add_table(rows=rows, cols=2)
    # setup table format:
    set_column_width(main_table.columns[0], Cm(4.05))
    set_column_width(main_table.columns[1], Cm(4.05))
    # Add descriptions to the first column of the main table:
    populate_description_columns(main_table, cif)
    # The main residuals table:
    populate_main_table_values(main_table, cif)
    return table_num


def make_report_from(files: List[Path], output_filename: str = 'tables.docx', path: str = '') -> str:
    nfiles = len(files)
    group_of_files = list(grouper(nfiles, 3))

    document = create_document(path)
    document.add_heading('Structure Tables', 1)

    for page_numbers in enumerate(group_of_files):

        main_table = document.add_table(rows=1, cols=4)
        populate_description_columns(main_table)
        
        for table_column in range(0, 3):  # the three columns
            if page_numbers[1][table_column]:
                cif_fileobj = files[table_column]
                filename = cif_fileobj.name
                cif = None
                try:
                    cif = CifContainer(cif_fileobj)
                except Exception as e:
                    print('Unable to open file', filename)
                    print(e)
                exti = cif['_refine_ls_extinction_coef']
                rows = 33
                if cif.is_centrosymm:
                    rows -= 1
                # Remove one row for the extinction coefficient:
                if exti in ['.', "'.'", '?', '']:
                    rows -= 1


    return ''
